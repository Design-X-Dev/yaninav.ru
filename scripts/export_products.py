#!/usr/bin/env python3
"""Экспорт товаров из products.json: папка на товар с фото, description.md и description.docx."""
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm
from PIL import Image, ImageOps

ROOT = Path(__file__).resolve().parents[1]
PRODUCTS_JSON = ROOT / "src" / "data" / "products.json"
IMAGES_DIR = ROOT / "public" / "images" / "products"
OUTPUT_DIR = ROOT / "output"

# Превью в папке товара и в DOCX: длинная сторона не больше этого значения (px)
PREVIEW_MAX_SIDE_PX = 1500
JPEG_PREVIEW_QUALITY = 85
WEBP_PREVIEW_QUALITY = 82

# Компактное отображение фото: до блока «Описание»
DOCX_IMAGE_WIDTH_CM = 7
MD_IMAGE_WIDTH_PX = 360


def format_price_ru(price: int | None) -> str:
    if price is None:
        return "По запросу"
    # Группировка тысяч пробелами
    s = f"{price:,}"
    return s.replace(",", " ") + " руб."


def slug_folder_name(product_id: int, name: str, max_len: int = 80) -> str:
    base = name.strip().lower()
    base = re.sub(r"\s+", "-", base)
    base = re.sub(r'[<>:"/\\|?*]', "", base)
    base = re.sub(r"-+", "-", base).strip("-")
    if not base:
        base = "tovar"
    if len(base) > max_len:
        base = base[:max_len].rstrip("-")
    return f"{product_id:02d}_{base}"


def collect_image_filenames(product: dict) -> list[str]:
    """Порядок: image, bannerImage, image2, image3; без пустых и дубликатов."""
    keys = ("image", "bannerImage", "image2", "image3")
    seen: set[str] = set()
    out: list[str] = []
    for key in keys:
        val = product.get(key)
        if not val or not isinstance(val, str):
            continue
        fname = val.strip()
        if not fname:
            continue
        if not fname.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
            continue
        if fname not in seen:
            seen.add(fname)
            out.append(fname)
    return out


def copy_product_images(
    filenames: list[str], dest_dir: Path
) -> tuple[list[str], list[str]]:
    """Копирует файлы в dest_dir. Возвращает (скопированные имена, отсутствующие в источнике)."""
    copied: list[str] = []
    missing: list[str] = []
    for fname in filenames:
        src = IMAGES_DIR / fname
        if not src.is_file():
            missing.append(fname)
            continue
        shutil.copy2(src, dest_dir / fname)
        copied.append(fname)
    return copied, missing


def compress_preview_image(path: Path, max_side: int = PREVIEW_MAX_SIDE_PX) -> None:
    """Уменьшает изображение: max(ширина, высота) <= max_side, пересохраняет со сжатием."""
    ext = path.suffix.lower()
    with Image.open(path) as im:
        im = ImageOps.exif_transpose(im)
        im.load()
        w, h = im.size
        if w < 1 or h < 1:
            return
        if max(w, h) > max_side:
            scale = max_side / float(max(w, h))
            new_w = max(1, int(round(w * scale)))
            new_h = max(1, int(round(h * scale)))
            im = im.resize((new_w, new_h), Image.Resampling.LANCZOS)

        if ext in (".jpg", ".jpeg"):
            rgb = im
            if rgb.mode in ("RGBA", "P"):
                rgb = rgb.convert("RGB")
            elif rgb.mode == "CMYK":
                rgb = rgb.convert("RGB")
            elif rgb.mode != "RGB":
                rgb = rgb.convert("RGB")
            rgb.save(
                path,
                format="JPEG",
                quality=JPEG_PREVIEW_QUALITY,
                optimize=True,
                progressive=True,
            )
        elif ext == ".png":
            if im.mode == "P" and "transparency" in im.info:
                im = im.convert("RGBA")
            im.save(path, format="PNG", optimize=True)
        elif ext == ".webp":
            im.save(
                path,
                format="WEBP",
                quality=WEBP_PREVIEW_QUALITY,
                method=6,
            )
        else:
            im.save(path)


def optimize_copied_images(folder: Path, filenames: list[str]) -> None:
    for fname in filenames:
        p = folder / fname
        if not p.is_file():
            continue
        try:
            compress_preview_image(p)
        except OSError as e:
            print(f"Предупреждение: не удалось обработать {p}: {e}")


def write_markdown(
    path: Path,
    name: str,
    category: str,
    price_line: str,
    description: str,
    image_files: list[str],
) -> None:
    lines = [
        f"# {name}",
        "",
        f"**Категория:** {category}",
        f"**Цена:** {price_line}",
        "",
        "## Фотографии",
        "",
    ]
    for i, fname in enumerate(image_files, start=1):
        lines.append(
            f'<img src="{fname}" alt="Фото {i}" width="{MD_IMAGE_WIDTH_PX}" />'
        )
        lines.append("")
    lines.extend(
        [
            "## Описание",
            "",
            description.strip(),
            "",
        ]
    )
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_docx(
    path: Path,
    name: str,
    category: str,
    price_line: str,
    description: str,
    image_dir: Path,
    image_files: list[str],
) -> None:
    doc = Document()
    doc.add_heading(name, level=1)

    p = doc.add_paragraph()
    p.add_run("Категория: ").bold = True
    p.add_run(category)

    p = doc.add_paragraph()
    p.add_run("Цена: ").bold = True
    p.add_run(price_line)

    doc.add_heading("Фотографии", level=2)
    for fname in image_files:
        img_path = image_dir / fname
        if img_path.is_file():
            doc.add_picture(str(img_path), width=Cm(DOCX_IMAGE_WIDTH_CM))
            doc.add_paragraph()

    doc.add_heading("Описание", level=2)
    for block in description.strip().split("\n"):
        doc.add_paragraph(block)

    doc.save(str(path))


def main() -> None:
    if not PRODUCTS_JSON.is_file():
        raise SystemExit(f"Не найден файл каталога: {PRODUCTS_JSON}")
    if not IMAGES_DIR.is_dir():
        raise SystemExit(f"Не найдена папка с фото: {IMAGES_DIR}")

    products = json.loads(PRODUCTS_JSON.read_text(encoding="utf-8"))
    if not isinstance(products, list):
        raise SystemExit("Ожидался JSON-массив товаров")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    total_copied = 0
    all_missing: list[tuple[int, str]] = []

    for product in products:
        pid = int(product["id"])
        name = str(product.get("name", "")).strip() or f"Товар {pid}"
        category = str(product.get("category", "")).strip()
        description = str(product.get("description", ""))
        price = product.get("price")
        if price is not None and not isinstance(price, int):
            try:
                price = int(price)
            except (TypeError, ValueError):
                price = None

        price_line = format_price_ru(price)
        folder = OUTPUT_DIR / slug_folder_name(pid, name)
        folder.mkdir(parents=True, exist_ok=True)

        image_files = collect_image_filenames(product)
        copied, missing = copy_product_images(image_files, folder)
        for m in missing:
            all_missing.append((pid, m))

        optimize_copied_images(folder, copied)

        write_markdown(
            folder / "description.md",
            name,
            category,
            price_line,
            description,
            copied,
        )
        write_docx(
            folder / "description.docx",
            name,
            category,
            price_line,
            description,
            folder,
            copied,
        )
        total_copied += len(copied)

    print(f"Обработано товаров: {len(products)}")
    print(f"Скопировано файлов изображений: {total_copied}")
    print(f"Выходная папка: {OUTPUT_DIR}")
    if all_missing:
        print("Предупреждение: отсутствуют файлы в public/images/products/:")
        for pid, fname in all_missing:
            print(f"  id={pid}: {fname}")


if __name__ == "__main__":
    main()
